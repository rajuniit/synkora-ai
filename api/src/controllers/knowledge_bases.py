"""Knowledge Base API endpoints."""

import logging
from uuid import UUID

import magic
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field
from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.database import get_async_db
from src.middleware.auth_middleware import get_current_tenant_id
from src.models.data_source import DataSource, DataSourceStatus, DataSourceType
from src.models.document import Document, DocumentStatus
from src.models.document_segment import DocumentSegment
from src.models.knowledge_base import (
    ChunkingStrategy,
    EmbeddingProvider,
    KnowledgeBase,
    KnowledgeBaseStatus,
    VectorDBProvider,
)
from src.models.tenant import Tenant, TenantPlan, TenantStatus
from src.services.knowledge_base import RAGService
from src.services.knowledge_base.embedding_service import EmbeddingService
from src.services.security.file_security import FileSecurityService, ScannerUnavailableError
from src.services.storage.s3_storage import S3StorageService
from src.tasks.kb_tasks import crawl_and_process_kb, process_kb_documents

# SECURITY: Global file security service instance for KB document uploads
_kb_file_security = FileSecurityService()

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/knowledge-bases", tags=["knowledge-bases"])


# Request/Response Models
class CreateKnowledgeBaseRequest(BaseModel):
    """Request model for creating a knowledge base."""

    name: str = Field(..., min_length=1, max_length=255)
    description: str | None = None
    embedding_provider: str = Field(default="SENTENCE_TRANSFORMERS")
    embedding_model: str = Field(default="all-MiniLM-L6-v2")
    embedding_config: dict = Field(default_factory=dict)
    vector_db_provider: str = Field(default="QDRANT")
    vector_db_config: dict = Field(default_factory=dict)
    chunking_strategy: str = Field(default="SEMANTIC")
    chunk_size: int = Field(default=1500, ge=100, le=10000)
    chunk_overlap: int = Field(default=150, ge=0, le=1000)
    min_chunk_size: int = Field(default=500, ge=100, le=5000)
    max_chunk_size: int = Field(default=3000, ge=1000, le=10000)
    chunking_config: dict = Field(default_factory=dict)


class UpdateKnowledgeBaseRequest(BaseModel):
    """Request model for updating a knowledge base."""

    name: str | None = Field(None, min_length=1, max_length=255)
    description: str | None = None
    embedding_provider: str | None = None
    embedding_model: str | None = None
    embedding_config: dict | None = None
    vector_db_provider: str | None = None
    vector_db_config: dict | None = None
    chunking_strategy: str | None = None
    chunk_size: int | None = Field(None, ge=100, le=10000)
    chunk_overlap: int | None = Field(None, ge=0, le=1000)
    min_chunk_size: int | None = Field(None, ge=100, le=5000)
    max_chunk_size: int | None = Field(None, ge=1000, le=10000)
    chunking_config: dict | None = None


class KnowledgeBaseResponse(BaseModel):
    """Response model for knowledge base."""

    id: int
    name: str
    description: str | None
    tenant_id: str
    embedding_provider: str
    embedding_model: str
    embedding_config: dict
    vector_db_provider: str
    vector_db_config: dict
    chunking_strategy: str
    chunk_size: int
    chunk_overlap: int
    min_chunk_size: int
    max_chunk_size: int
    chunking_config: dict
    document_count: int
    total_chunks: int
    is_active: bool
    created_at: str
    updated_at: str

    model_config = ConfigDict(from_attributes=True)


class KnowledgeBaseStatsResponse(BaseModel):
    """Response model for knowledge base statistics."""

    id: int
    name: str
    document_count: int
    total_chunks: int
    data_sources: list[dict]
    last_sync: str | None
    storage_size_mb: float | None


class SearchRequest(BaseModel):
    """Request model for searching knowledge base."""

    query: str = Field(..., min_length=1)
    limit: int = Field(default=5, ge=1, le=50)
    min_score: float = Field(default=0.7, ge=0.0, le=1.0)


class SearchResult(BaseModel):
    """Search result model."""

    document_id: int
    chunk_id: int
    text: str
    score: float
    metadata: dict


# Endpoints
async def get_default_tenant(db: AsyncSession) -> UUID:
    """Get or create default tenant."""
    result = await db.execute(select(Tenant).limit(1))
    tenant = result.scalar_one_or_none()
    if not tenant:
        # Create default tenant
        tenant = Tenant(name="Default Tenant", plan=TenantPlan.FREE, status=TenantStatus.ACTIVE)
        db.add(tenant)
        await db.commit()
        await db.refresh(tenant)
    return tenant.id


@router.post("", response_model=KnowledgeBaseResponse, status_code=201)
async def create_knowledge_base(
    request: CreateKnowledgeBaseRequest,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Create a new knowledge base."""
    try:
        # Convert string values to enums
        try:
            embedding_provider_enum = EmbeddingProvider(request.embedding_provider)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid embedding provider: {request.embedding_provider}")

        try:
            vector_db_provider_enum = VectorDBProvider(request.vector_db_provider)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid vector DB provider: {request.vector_db_provider}")

        try:
            chunking_strategy_enum = ChunkingStrategy(request.chunking_strategy)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid chunking strategy: {request.chunking_strategy}")

        kb = KnowledgeBase(
            name=request.name,
            description=request.description,
            tenant_id=tenant_id,
            embedding_provider=embedding_provider_enum,
            embedding_model=request.embedding_model,
            vector_db_provider=vector_db_provider_enum,
            chunking_strategy=chunking_strategy_enum,
            chunk_size=request.chunk_size,
            chunk_overlap=request.chunk_overlap,
            min_chunk_size=request.min_chunk_size,
            max_chunk_size=request.max_chunk_size,
            chunking_config=request.chunking_config,
            status=KnowledgeBaseStatus.ACTIVE,
        )

        # Set configs with encryption
        kb.set_embedding_config_encrypted(request.embedding_config)
        kb.set_vector_db_config_encrypted(request.vector_db_config)

        db.add(kb)
        await db.commit()
        await db.refresh(kb)

        logger.info(f"Created knowledge base: {kb.name} (ID: {kb.id})")

        return KnowledgeBaseResponse(
            id=kb.id,
            name=kb.name,
            description=kb.description,
            tenant_id=str(kb.tenant_id),
            embedding_provider=kb.embedding_provider,
            embedding_model=kb.embedding_model,
            embedding_config=kb.get_embedding_config_decrypted(),
            vector_db_provider=kb.vector_db_provider,
            vector_db_config=kb.get_vector_db_config_decrypted(),
            chunking_strategy=kb.chunking_strategy,
            chunk_size=kb.chunk_size,
            chunk_overlap=kb.chunk_overlap,
            min_chunk_size=kb.min_chunk_size,
            max_chunk_size=kb.max_chunk_size,
            chunking_config=kb.chunking_config or {},
            document_count=kb.total_documents,
            total_chunks=kb.total_chunks,
            is_active=(kb.status == KnowledgeBaseStatus.ACTIVE),
            created_at=kb.created_at.isoformat(),
            updated_at=kb.updated_at.isoformat(),
        )

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        logger.error(f"Error creating knowledge base: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("", response_model=list[KnowledgeBaseResponse])
async def list_knowledge_bases(
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=100),
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """List all knowledge bases."""
    try:
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.tenant_id == tenant_id).offset(skip).limit(limit)
        )
        kbs = result.scalars().all()

        return [
            KnowledgeBaseResponse(
                id=kb.id,
                name=kb.name,
                description=kb.description,
                tenant_id=str(kb.tenant_id),
                embedding_provider=kb.embedding_provider,
                embedding_model=kb.embedding_model,
                embedding_config=kb.embedding_config or {},
                vector_db_provider=kb.vector_db_provider,
                vector_db_config=kb.vector_db_config or {},
                chunking_strategy=kb.chunking_strategy,
                chunk_size=kb.chunk_size,
                chunk_overlap=kb.chunk_overlap,
                min_chunk_size=kb.min_chunk_size,
                max_chunk_size=kb.max_chunk_size,
                chunking_config=kb.chunking_config or {},
                document_count=kb.total_documents,
                total_chunks=kb.total_chunks,
                is_active=(kb.status == KnowledgeBaseStatus.ACTIVE),
                created_at=kb.created_at.isoformat(),
                updated_at=kb.updated_at.isoformat(),
            )
            for kb in kbs
        ]

    except Exception as e:
        logger.error(f"Error listing knowledge bases: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{kb_id}", response_model=KnowledgeBaseResponse)
async def get_knowledge_base(
    kb_id: int,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Get a specific knowledge base.

    SECURITY: Requires authentication and verifies KB belongs to tenant.
    """
    try:
        # SECURITY: Filter by tenant_id to prevent IDOR attacks
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        return KnowledgeBaseResponse(
            id=kb.id,
            name=kb.name,
            description=kb.description,
            tenant_id=str(kb.tenant_id),
            embedding_provider=kb.embedding_provider,
            embedding_model=kb.embedding_model,
            embedding_config=kb.embedding_config or {},
            vector_db_provider=kb.vector_db_provider,
            vector_db_config=kb.vector_db_config or {},
            chunking_strategy=kb.chunking_strategy,
            chunk_size=kb.chunk_size,
            chunk_overlap=kb.chunk_overlap,
            min_chunk_size=kb.min_chunk_size,
            max_chunk_size=kb.max_chunk_size,
            chunking_config=kb.chunking_config or {},
            document_count=kb.total_documents,
            total_chunks=kb.total_chunks,
            is_active=(kb.status == KnowledgeBaseStatus.ACTIVE),
            created_at=kb.created_at.isoformat(),
            updated_at=kb.updated_at.isoformat(),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting knowledge base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/{kb_id}", response_model=KnowledgeBaseResponse)
async def update_knowledge_base(
    kb_id: int,
    request: UpdateKnowledgeBaseRequest,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Update a knowledge base.

    SECURITY: Requires authentication and verifies KB belongs to tenant.
    """
    try:
        # SECURITY: Filter by tenant_id to prevent IDOR attacks
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Update fields
        if request.name is not None:
            kb.name = request.name
        if request.description is not None:
            kb.description = request.description
        if request.embedding_provider is not None:
            kb.embedding_provider = request.embedding_provider
        if request.embedding_model is not None:
            kb.embedding_model = request.embedding_model
        if request.embedding_config is not None:
            kb.set_embedding_config_encrypted(request.embedding_config)
        if request.vector_db_provider is not None:
            kb.vector_db_provider = request.vector_db_provider
        if request.vector_db_config is not None:
            kb.set_vector_db_config_encrypted(request.vector_db_config)
        if request.chunking_strategy is not None:
            kb.chunking_strategy = request.chunking_strategy
        if request.chunk_size is not None:
            kb.chunk_size = request.chunk_size
        if request.chunk_overlap is not None:
            kb.chunk_overlap = request.chunk_overlap
        if request.min_chunk_size is not None:
            kb.min_chunk_size = request.min_chunk_size
        if request.max_chunk_size is not None:
            kb.max_chunk_size = request.max_chunk_size
        if request.chunking_config is not None:
            kb.chunking_config = request.chunking_config

        await db.commit()
        await db.refresh(kb)

        logger.info(f"Updated knowledge base: {kb.name} (ID: {kb.id})")

        return KnowledgeBaseResponse(
            id=kb.id,
            name=kb.name,
            description=kb.description,
            tenant_id=str(kb.tenant_id),
            embedding_provider=kb.embedding_provider,
            embedding_model=kb.embedding_model,
            embedding_config=kb.embedding_config or {},
            vector_db_provider=kb.vector_db_provider,
            vector_db_config=kb.vector_db_config or {},
            chunking_strategy=kb.chunking_strategy,
            chunk_size=kb.chunk_size,
            chunk_overlap=kb.chunk_overlap,
            min_chunk_size=kb.min_chunk_size,
            max_chunk_size=kb.max_chunk_size,
            chunking_config=kb.chunking_config or {},
            document_count=kb.total_documents,
            total_chunks=kb.total_chunks,
            is_active=(kb.status == KnowledgeBaseStatus.ACTIVE),
            created_at=kb.created_at.isoformat(),
            updated_at=kb.updated_at.isoformat(),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating knowledge base: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{kb_id}", status_code=204)
async def delete_knowledge_base(
    kb_id: int,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Delete a knowledge base.

    SECURITY: Requires authentication and verifies KB belongs to tenant.
    """
    try:
        # SECURITY: Filter by tenant_id to prevent IDOR attacks
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Delete associated documents
        await db.execute(delete(Document).filter(Document.knowledge_base_id == kb_id))

        # Delete knowledge base
        await db.delete(kb)
        await db.commit()

        logger.info(f"Deleted knowledge base: {kb.name} (ID: {kb.id})")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting knowledge base: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{kb_id}/stats", response_model=KnowledgeBaseStatsResponse)
async def get_knowledge_base_stats(
    kb_id: int,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Get statistics for a knowledge base.

    SECURITY: Requires authentication and verifies KB belongs to tenant.
    """
    try:
        # SECURITY: Filter by tenant_id to prevent IDOR attacks
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Get data sources (limit to prevent loading thousands of rows)
        result = await db.execute(select(DataSource).filter(DataSource.knowledge_base_id == kb_id).limit(500))
        data_sources = result.scalars().all()

        data_source_info = [
            {
                "id": ds.id,
                "name": ds.name,
                "type": ds.type.value,
                "is_connected": ds.is_connected,
                "last_sync": ds.last_sync_at.isoformat() if ds.last_sync_at else None,
            }
            for ds in data_sources
        ]

        # Get last sync time
        last_sync = None
        if data_sources:
            sync_times = [ds.last_sync_at for ds in data_sources if ds.last_sync_at]
            if sync_times:
                last_sync = max(sync_times).isoformat()

        return KnowledgeBaseStatsResponse(
            id=kb.id,
            name=kb.name,
            document_count=kb.total_documents,
            total_chunks=kb.total_chunks,
            data_sources=data_source_info,
            last_sync=last_sync,
            storage_size_mb=None,  # Vector DB storage tracking not implemented
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting knowledge base stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{kb_id}/search", response_model=list[SearchResult])
async def search_knowledge_base(
    kb_id: int,
    request: SearchRequest,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Search a knowledge base.

    SECURITY: Requires authentication and verifies KB belongs to tenant.
    """
    try:
        # SECURITY: Filter by tenant_id to prevent IDOR attacks
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Initialize services
        embedding_service = EmbeddingService()
        rag_service = RAGService(db, embedding_service.client)

        # Search
        results = await rag_service.search_knowledge_base(
            knowledge_base=kb, query=request.query, limit=request.limit, min_score=request.min_score
        )

        return [
            SearchResult(
                document_id=r["document_id"],
                chunk_id=r["chunk_id"],
                text=r["text"],
                score=r["score"],
                metadata=r["metadata"],
            )
            for r in results
        ]

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error searching knowledge base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class DocumentUploadResponse(BaseModel):
    """Response model for document upload."""

    success: bool
    message: str
    documents_processed: int
    documents_embedded: int
    total_chunks: int
    failed_files: list[dict] = []


class AddTextContentRequest(BaseModel):
    """Request model for adding text content directly."""

    title: str = Field(..., min_length=1, max_length=255)
    content: str = Field(..., min_length=1)
    metadata: dict = Field(default_factory=dict)


class CrawlWebsiteRequest(BaseModel):
    """Request model for crawling a website."""

    url: str = Field(..., min_length=1)
    max_pages: int = Field(default=1, ge=1, le=50)
    include_subpages: bool = Field(default=False)


# File upload constants
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MAX_FILES_PER_UPLOAD = 10
ALLOWED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/html": "html",
    "text/csv": "csv",
}


@router.post("/{kb_id}/documents/upload", response_model=DocumentUploadResponse)
async def upload_documents(
    kb_id: int,
    files: list[UploadFile] = File(...),
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """
    Upload documents directly to a knowledge base.

    Supports: PDF, DOCX, TXT, MD, HTML, CSV
    Max file size: 50MB per file
    Max files: 10 per request
    """
    try:
        # Get knowledge base
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Validate number of files
        if len(files) > MAX_FILES_PER_UPLOAD:
            raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES_PER_UPLOAD} files allowed per upload")

        # Initialize services
        s3_storage = S3StorageService()

        failed_files = []
        processed_docs = []

        for file in files:
            try:
                # Read file content
                content = await file.read()
                file_size = len(content)

                # Validate file size
                if file_size > MAX_FILE_SIZE:
                    failed_files.append(
                        {
                            "filename": file.filename,
                            "error": f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds maximum ({MAX_FILE_SIZE / 1024 / 1024}MB)",
                        }
                    )
                    continue

                # SECURITY: Comprehensive file validation (magic bytes, dangerous extensions,
                # malicious content patterns) before any further processing.
                try:
                    validation_result = _kb_file_security.validate_file(
                        file_content=content,
                        filename=file.filename or "unnamed",
                        category="document",
                    )
                    if not validation_result["is_valid"]:
                        error_details = "; ".join(validation_result["errors"])
                        logger.warning(
                            f"KB document upload rejected for tenant {tenant_id}: {file.filename} — {error_details}"
                        )
                        failed_files.append(
                            {"filename": file.filename, "error": f"File validation failed: {error_details}"}
                        )
                        continue
                except ScannerUnavailableError as scan_err:
                    logger.error(f"File scanning service unavailable: {scan_err}")
                    raise HTTPException(
                        status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                        detail="File scanning service unavailable. Please try again later.",
                    )

                # Detect MIME type
                mime = magic.from_buffer(content, mime=True)

                # Validate file type
                if mime not in ALLOWED_MIME_TYPES:
                    failed_files.append({"filename": file.filename, "error": f"Unsupported file type: {mime}"})
                    continue

                file_type = ALLOWED_MIME_TYPES[mime]

                # SECURITY: Sanitize filename to prevent path traversal attacks
                import re
                import uuid as uuid_module

                # Remove any path components and dangerous characters
                safe_filename = re.sub(r"[^\w\s\-\.]", "", file.filename.replace("/", "_").replace("\\", "_"))
                safe_filename = safe_filename.strip(". ")  # Remove leading/trailing dots and spaces
                if not safe_filename:
                    safe_filename = f"file_{uuid_module.uuid4().hex[:8]}"
                # Add unique prefix to prevent overwrites
                unique_filename = f"{uuid_module.uuid4().hex[:8]}_{safe_filename}"

                # Generate S3 key with sanitized filename
                s3_key = f"knowledge-bases/kb-{kb.id}/uploads/{unique_filename}"

                # Upload to S3
                logger.info(f"Uploading {file.filename} to S3...")
                s3_result = s3_storage.upload_file(
                    file_content=content,
                    key=s3_key,
                    content_type=mime,
                    metadata={
                        "knowledge_base_id": str(kb.id),
                        "tenant_id": str(tenant_id),
                        "original_filename": file.filename,
                        "upload_source": "ui",
                    },
                )

                # Extract text content based on file type
                text_content = ""
                if file_type == "txt" or file_type == "md" or file_type == "csv":
                    text_content = content.decode("utf-8", errors="ignore")
                elif file_type == "html":
                    from bs4 import BeautifulSoup

                    soup = BeautifulSoup(content, "html.parser")
                    text_content = soup.get_text(separator="\n", strip=True)
                elif file_type == "pdf":
                    import fitz  # PyMuPDF

                    doc = fitz.open(stream=content, filetype="pdf")
                    text_content = "\n".join([page.get_text() for page in doc])
                    doc.close()
                elif file_type == "docx":
                    from io import BytesIO

                    from docx import Document as DocxDocument

                    doc = DocxDocument(BytesIO(content))
                    text_content = "\n".join([para.text for para in doc.paragraphs])

                # Create a stub Document record immediately so the file
                # is visible in the browser before Celery finishes.
                stub_doc = Document(
                    tenant_id=tenant_id,
                    knowledge_base_id=kb.id,
                    name=file.filename,
                    external_id=file.filename,
                    s3_key=s3_key,
                    s3_url=s3_result.get("url"),
                    file_size=file_size,
                    mime_type=mime,
                    original_filename=file.filename,
                    source_type=DataSourceType.MANUAL.value,
                    content_type=file_type,
                    upload_source="ui",
                    status=DocumentStatus.PENDING,
                    doc_metadata={
                        "title": file.filename,
                        "file_type": file_type,
                        "upload_source": "ui",
                    },
                )
                db.add(stub_doc)
                await db.flush()  # get stub_doc.id without committing yet

                processed_docs.append(
                    {
                        "id": file.filename,
                        "text": text_content,
                        "metadata": {
                            "title": file.filename,
                            "url": s3_result.get("url", ""),
                            "s3_url": s3_result.get("url", ""),
                            "file_type": file_type,
                            "file_size": file_size,
                            "mime_type": mime,
                            "original_filename": file.filename,
                            "upload_source": "ui",
                        },
                    }
                )

                logger.info(f"✓ Processed {file.filename}")

            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {e}")
                failed_files.append({"filename": file.filename, "error": str(e)})
                continue

        # Dispatch chunking + embedding to Celery so the HTTP request returns immediately
        if processed_docs:
            # Get or create the manual data source
            result = await db.execute(
                select(DataSource).filter(
                    DataSource.knowledge_base_id == kb.id, DataSource.type == DataSourceType.MANUAL
                )
            )
            data_source = result.scalar_one_or_none()

            if not data_source:
                data_source = DataSource(
                    tenant_id=tenant_id,
                    knowledge_base_id=kb.id,
                    name="Manual Uploads",
                    type=DataSourceType.MANUAL,
                    status=DataSourceStatus.ACTIVE,
                )
                db.add(data_source)

            # Update KB document count to include the stub docs
            kb.total_documents = (kb.total_documents or 0) + len(processed_docs)

            # Commit stubs + data_source together so docs are immediately visible
            await db.commit()
            await db.refresh(data_source)

            process_kb_documents.delay(data_source.id, str(tenant_id), processed_docs)

            return DocumentUploadResponse(
                success=True,
                message=f"{len(processed_docs)} document(s) uploaded and queued for processing",
                documents_processed=len(processed_docs),
                documents_embedded=0,
                total_chunks=0,
                failed_files=failed_files,
            )
        else:
            return DocumentUploadResponse(
                success=False,
                message="No documents were successfully processed",
                documents_processed=0,
                documents_embedded=0,
                total_chunks=0,
                failed_files=failed_files,
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading documents: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{kb_id}/documents/text", response_model=DocumentUploadResponse)
async def add_text_content(
    kb_id: int,
    request: AddTextContentRequest,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """
    Add text content directly to a knowledge base.

    This endpoint allows users to paste text content without uploading a file.
    """
    try:
        # Get knowledge base
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Get or create manual data source
        result = await db.execute(
            select(DataSource).filter(DataSource.knowledge_base_id == kb.id, DataSource.type == DataSourceType.MANUAL)
        )
        data_source = result.scalar_one_or_none()

        if not data_source:
            data_source = DataSource(
                tenant_id=tenant_id,
                knowledge_base_id=kb.id,
                name="Manual Content",
                type=DataSourceType.MANUAL,
                status=DataSourceStatus.ACTIVE,
            )
            db.add(data_source)
            await db.commit()
            await db.refresh(data_source)

        documents = [
            {
                "id": request.title,
                "text": request.content,
                "metadata": {
                    "title": request.title,
                    "source_type": "text",
                    "upload_source": "paste",
                    **request.metadata,
                },
            }
        ]

        # Dispatch chunking + embedding to Celery
        process_kb_documents.delay(data_source.id, str(tenant_id), documents)

        logger.info(f"Queued text content '{request.title}' for KB {kb_id}")

        return DocumentUploadResponse(
            success=True,
            message=f"Text content '{request.title}' queued for processing",
            documents_processed=1,
            documents_embedded=0,
            total_chunks=0,
            failed_files=[],
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error adding text content: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{kb_id}/documents/crawl", response_model=DocumentUploadResponse)
async def crawl_website(
    kb_id: int,
    request: CrawlWebsiteRequest,
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """
    Crawl a website URL and add its content to the knowledge base.

    This endpoint fetches content from a URL and processes it.
    """
    try:
        from urllib.parse import urlparse

        from src.services.security.url_validator import validate_url

        # Get knowledge base
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Validate URL format and SSRF before queuing
        parsed_url = urlparse(request.url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise HTTPException(status_code=400, detail="Invalid URL format")

        is_valid, error_msg = validate_url(
            request.url, allowed_schemes=["http", "https"], block_private_ips=True, resolve_dns=True
        )
        if not is_valid:
            raise HTTPException(status_code=400, detail=f"URL not allowed: {error_msg}")

        # Get or create web data source
        result = await db.execute(
            select(DataSource).filter(DataSource.knowledge_base_id == kb.id, DataSource.type == DataSourceType.WEB)
        )
        data_source = result.scalar_one_or_none()

        if not data_source:
            data_source = DataSource(
                tenant_id=tenant_id,
                knowledge_base_id=kb.id,
                name="Web Content",
                type=DataSourceType.WEB,
                status=DataSourceStatus.ACTIVE,
            )
            db.add(data_source)
            await db.commit()
            await db.refresh(data_source)

        # Dispatch crawl + embed to Celery
        crawl_and_process_kb.delay(
            data_source.id,
            str(tenant_id),
            request.url,
            request.max_pages,
            request.include_subpages,
        )

        logger.info(f"Queued crawl of {request.url} for KB {kb_id}")

        return DocumentUploadResponse(
            success=True,
            message=f"Crawl of '{request.url}' queued for processing",
            documents_processed=0,
            documents_embedded=0,
            total_chunks=0,
            failed_files=[],
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error queuing website crawl: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# Document Browser Models
class DocumentResponse(BaseModel):
    """Response model for document."""

    id: str
    title: str
    source_type: str
    source_url: str | None
    file_size: int | None
    chunk_count: int
    has_images: bool
    image_count: int
    status: str
    created_at: str
    updated_at: str
    metadata: dict

    model_config = ConfigDict(from_attributes=True)


class DocumentListResponse(BaseModel):
    """Response model for document list with pagination."""

    documents: list[DocumentResponse]
    total: int
    page: int
    page_size: int
    total_pages: int


class DocumentDetailResponse(BaseModel):
    """Response model for detailed document view."""

    id: str
    title: str
    source_type: str
    source_url: str | None
    s3_url: str | None
    mime_type: str | None
    file_size: int | None
    chunk_count: int
    has_images: bool
    image_count: int
    images: list[dict]
    created_at: str
    updated_at: str
    metadata: dict
    data_source: dict | None
    content: str | None  # Full document content
    segments: list[dict]  # List of chunks/segments

    model_config = ConfigDict(from_attributes=True)


@router.get("/{kb_id}/documents", response_model=DocumentListResponse)
async def list_documents(
    kb_id: int,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    search: str | None = Query(None),
    source_type: str | None = Query(None),
    has_images: bool | None = Query(None),
    sort_by: str = Query("created_at", pattern="^(created_at|updated_at|title|chunk_count)$"),
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """
    List all documents in a knowledge base with pagination and filtering.

    Query Parameters:
    - page: Page number (default: 1)
    - page_size: Items per page (default: 20, max: 100)
    - search: Search in document titles
    - source_type: Filter by source type (manual, SLACK, gmail, etc.)
    - has_images: Filter documents with/without images
    - sort_by: Sort field (created_at, updated_at, title, chunk_count)
    - sort_order: Sort order (asc, desc)
    """
    try:
        # Verify knowledge base exists and belongs to tenant
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Build query
        query = select(Document).filter(Document.knowledge_base_id == kb_id)

        # Apply filters
        if search:
            # SECURITY: Escape special LIKE pattern characters to prevent SQL injection
            # The characters % and _ have special meaning in LIKE patterns
            escaped_search = search.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
            query = query.filter(Document.name.ilike(f"%{escaped_search}%", escape="\\"))

        if source_type:
            query = query.filter(Document.source_type == source_type)

        if has_images is not None:
            query = query.filter(Document.has_images == has_images)

        # Get total count
        count_result = await db.execute(select(func.count()).select_from(query.subquery()))
        total = count_result.scalar()

        # Apply sorting
        sort_column = getattr(Document, sort_by)
        if sort_order == "desc":
            query = query.order_by(sort_column.desc())
        else:
            query = query.order_by(sort_column.asc())

        # Apply pagination
        offset = (page - 1) * page_size
        result = await db.execute(query.offset(offset).limit(page_size))
        documents = result.scalars().all()

        # Batch-fetch segment counts — avoids lazy-loading segments relationship per doc
        doc_ids = [doc.id for doc in documents]
        segment_counts: dict = {}
        if doc_ids:
            seg_result = await db.execute(
                select(DocumentSegment.document_id, func.count(DocumentSegment.id))
                .filter(DocumentSegment.document_id.in_(doc_ids))
                .group_by(DocumentSegment.document_id)
            )
            segment_counts = dict(seg_result.all())

        # Calculate total pages
        total_pages = (total + page_size - 1) // page_size

        return DocumentListResponse(
            documents=[
                DocumentResponse(
                    id=str(doc.id),
                    title=doc.name,
                    source_type=doc.source_type,
                    source_url=doc.external_url,
                    file_size=doc.file_size,
                    chunk_count=segment_counts.get(doc.id, 0),
                    has_images=doc.has_images,
                    image_count=doc.image_count,
                    status=doc.status.value,
                    created_at=doc.created_at.isoformat(),
                    updated_at=doc.updated_at.isoformat(),
                    metadata=doc.doc_metadata or {},
                )
                for doc in documents
            ],
            total=total,
            page=page,
            page_size=page_size,
            total_pages=total_pages,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{kb_id}/documents/{doc_id}", response_model=DocumentDetailResponse)
async def get_document_details(
    kb_id: int, doc_id: str, tenant_id: UUID = Depends(get_current_tenant_id), db: AsyncSession = Depends(get_async_db)
):
    """Get detailed information about a specific document."""
    try:
        # Verify knowledge base exists and belongs to tenant
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Get document
        result = await db.execute(select(Document).filter(Document.id == doc_id, Document.knowledge_base_id == kb_id))
        doc = result.scalar_one_or_none()

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Get data source info if available
        data_source_info = None
        if doc.data_source_id:
            result = await db.execute(select(DataSource).filter(DataSource.id == doc.data_source_id))
            data_source = result.scalar_one_or_none()
            if data_source:
                data_source_info = {"id": data_source.id, "name": data_source.name, "type": data_source.type.value}

        # Get document segments (chunks) — limit to prevent loading entire large document into memory
        result = await db.execute(
            select(DocumentSegment)
            .filter(DocumentSegment.document_id == doc.id)
            .order_by(DocumentSegment.position)
            .limit(200)
        )
        segments = result.scalars().all()

        segments_data = [
            {
                "id": str(seg.id),
                "position": seg.position,
                "content": seg.content,
                "word_count": seg.word_count,
                "tokens": seg.tokens,
                "created_at": seg.created_at.isoformat(),
            }
            for seg in segments
        ]

        # Generate pre-signed URL - check both s3_url and external_url (source_url) fields
        s3_uri = doc.s3_url or doc.external_url
        presigned_url = None

        # If we have an S3 URI, generate pre-signed URL
        if s3_uri and s3_uri.startswith("s3://"):
            try:
                s3_storage = S3StorageService()
                presigned_url = s3_storage.generate_presigned_url(
                    s3_uri,
                    expiration=3600,  # 1 hour
                )
                logger.info(f"Generated presigned URL for document {doc.id}: {presigned_url[:50]}...")
            except Exception as e:
                logger.error(f"Failed to generate presigned URL for document {doc.id}: {e}")
                raise HTTPException(status_code=500, detail=f"Failed to generate presigned URL: {str(e)}")

        return DocumentDetailResponse(
            id=str(doc.id),
            title=doc.name,
            source_type=doc.source_type,
            source_url=doc.external_url if not (doc.external_url and doc.external_url.startswith("s3://")) else None,
            s3_url=presigned_url,  # Always use presigned URL
            mime_type=doc.mime_type,
            file_size=doc.file_size,
            chunk_count=len(segments),
            has_images=doc.has_images,
            image_count=doc.image_count,
            images=doc.images or [],
            created_at=doc.created_at.isoformat(),
            updated_at=doc.updated_at.isoformat(),
            metadata=doc.doc_metadata or {},
            data_source=data_source_info,
            content=doc.content,  # Full document content
            segments=segments_data,  # List of chunks
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document details: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{kb_id}/documents/{doc_id}", status_code=204)
async def delete_document(
    kb_id: int, doc_id: str, tenant_id: UUID = Depends(get_current_tenant_id), db: AsyncSession = Depends(get_async_db)
):
    """Delete a document from the knowledge base."""
    try:
        # Verify knowledge base exists and belongs to tenant
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Get document
        result = await db.execute(select(Document).filter(Document.id == doc_id, Document.knowledge_base_id == kb_id))
        doc = result.scalar_one_or_none()

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Note: Vector DB segments are cleaned up by the RAG service
        # S3 cleanup happens via storage service cascade

        # Delete document
        await db.delete(doc)
        await db.commit()

        logger.info(f"Deleted document: {doc.name} (ID: {doc.id})")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{kb_id}/documents/bulk-delete", status_code=204)
async def bulk_delete_documents(
    kb_id: int,
    document_ids: list[str],
    tenant_id: UUID = Depends(get_current_tenant_id),
    db: AsyncSession = Depends(get_async_db),
):
    """Delete multiple documents from the knowledge base."""
    try:
        # Verify knowledge base exists and belongs to tenant
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Convert string UUIDs to UUID objects
        uuid_ids = [UUID(doc_id) for doc_id in document_ids]

        # Delete documents
        result = await db.execute(
            delete(Document).filter(Document.id.in_(uuid_ids), Document.knowledge_base_id == kb_id)
        )
        deleted_count = result.rowcount

        await db.commit()

        logger.info(f"Bulk deleted {deleted_count} documents from KB {kb_id}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error bulk deleting documents: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{kb_id}/documents/{doc_id}/download")
async def download_document(
    kb_id: int, doc_id: str, tenant_id: UUID = Depends(get_current_tenant_id), db: AsyncSession = Depends(get_async_db)
):
    """
    Download or access the original document.

    For uploaded files: Downloads from S3
    For synced documents (Gmail, Slack, etc.): Redirects to external URL or returns content as text
    """
    try:
        # Verify knowledge base exists and belongs to tenant
        result = await db.execute(
            select(KnowledgeBase).filter(KnowledgeBase.id == kb_id, KnowledgeBase.tenant_id == tenant_id)
        )
        kb = result.scalar_one_or_none()

        if not kb:
            raise HTTPException(status_code=404, detail="Knowledge base not found")

        # Get document
        result = await db.execute(select(Document).filter(Document.id == doc_id, Document.knowledge_base_id == kb_id))
        doc = result.scalar_one_or_none()

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Check if document has S3 URL (uploaded files)
        if doc.s3_url:
            # Download from S3
            s3_storage = S3StorageService()

            try:
                file_content = await s3_storage.download_file_content(doc.s3_url)

                # Determine content type
                content_type = "application/octet-stream"
                if doc.mime_type:
                    content_type = doc.mime_type
                elif doc.doc_metadata and "file_type" in doc.doc_metadata:
                    file_type = doc.doc_metadata["file_type"]
                    if file_type in ALLOWED_MIME_TYPES.values():
                        # Reverse lookup MIME type
                        for mime, ftype in ALLOWED_MIME_TYPES.items():
                            if ftype == file_type:
                                content_type = mime
                                break

                # Return file as streaming response
                from io import BytesIO
                from urllib.parse import quote

                # Encode filename for Content-Disposition header (RFC 5987)
                encoded_filename = quote(doc.name)

                return StreamingResponse(
                    BytesIO(file_content),
                    media_type=content_type,
                    headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
                )

            except Exception as e:
                logger.error(f"Error downloading file from S3: {e}")
                raise HTTPException(status_code=500, detail="Failed to download document file")

        # For synced documents (Gmail, Slack, etc.)
        elif doc.external_url:
            # Validate scheme before redirecting to prevent open redirect to javascript:/data: URIs
            from urllib.parse import urlparse

            parsed = urlparse(doc.external_url)
            if parsed.scheme not in ("http", "https"):
                raise HTTPException(status_code=400, detail="Invalid document URL")

            from fastapi.responses import RedirectResponse

            return RedirectResponse(url=doc.external_url)

        # If no S3 URL or external URL, return content as text file
        elif doc.content:
            from io import BytesIO
            from urllib.parse import quote

            content_bytes = doc.content.encode("utf-8")

            # Encode filename for Content-Disposition header (RFC 5987)
            encoded_filename = quote(f"{doc.name}.txt")

            return StreamingResponse(
                BytesIO(content_bytes),
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
            )

        else:
            raise HTTPException(status_code=404, detail="Document content not available for download")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))
