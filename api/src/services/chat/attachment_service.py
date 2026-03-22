"""Chat attachment service for handling file uploads in chat."""

import io
import logging
from datetime import UTC, datetime
from typing import BinaryIO
from uuid import UUID, uuid4

import docx
import PyPDF2

from src.services.storage.s3_storage import S3StorageService

logger = logging.getLogger(__name__)

# Supported file types
SUPPORTED_IMAGE_TYPES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/gif",
    "image/webp",
}

SUPPORTED_DOCUMENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
    "text/plain",
    "text/markdown",
    "text/csv",
}

SUPPORTED_CODE_TYPES = {
    "text/javascript",
    "text/x-python",
    "text/x-java",
    "text/x-c",
    "text/x-c++",
    "text/x-go",
    "text/x-rust",
    "text/x-typescript",
    "application/json",
    "application/xml",
    "text/html",
    "text/css",
}

ALL_SUPPORTED_TYPES = SUPPORTED_IMAGE_TYPES | SUPPORTED_DOCUMENT_TYPES | SUPPORTED_CODE_TYPES

# File size limits
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILES_PER_MESSAGE = 5


class AttachmentService:
    """Service for handling chat file attachments."""

    def __init__(self):
        """Initialize attachment service."""
        self.s3_storage = S3StorageService()

    async def upload_attachment(
        self,
        file_content: bytes | BinaryIO,
        filename: str,
        content_type: str,
        tenant_id: UUID,
        conversation_id: UUID,
    ) -> dict:
        """
        Upload a file attachment for chat.

        Args:
            file_content: File content as bytes or file-like object
            filename: Original filename
            content_type: MIME type of the file
            tenant_id: Tenant UUID
            conversation_id: Conversation UUID

        Returns:
            dict: Attachment metadata

        Raises:
            ValueError: If file type not supported or file too large
        """
        # SECURITY: Validate actual file content type using magic bytes, not just client-provided header
        # This prevents content-type spoofing attacks (e.g., uploading executable as image/png)
        import magic

        if isinstance(file_content, bytes):
            actual_content_type = magic.from_buffer(file_content, mime=True)
        else:
            # Read first 2048 bytes for magic detection
            file_content.seek(0)
            header = file_content.read(2048)
            file_content.seek(0)
            actual_content_type = magic.from_buffer(header, mime=True)

        # Use the actual detected content type, not the client-provided one
        if actual_content_type not in ALL_SUPPORTED_TYPES:
            raise ValueError(
                f"Unsupported file type: {actual_content_type}. "
                f"Supported types: images (JPEG, PNG, GIF, WebP), "
                f"documents (PDF, DOCX, TXT, MD, CSV), and code files."
            )

        # Update content_type to the verified one
        content_type = actual_content_type

        # Validate file size
        if isinstance(file_content, bytes):
            file_size = len(file_content)
        else:
            # For file-like objects, seek to end to get size
            file_content.seek(0, 2)
            file_size = file_content.tell()
            file_content.seek(0)

        if file_size > MAX_FILE_SIZE:
            raise ValueError(f"File size ({file_size} bytes) exceeds maximum allowed size ({MAX_FILE_SIZE} bytes)")

        # Generate unique file ID
        file_id = str(uuid4())

        # Generate S3 key
        timestamp = datetime.now(UTC)
        date_path = timestamp.strftime("%Y/%m/%d")
        s3_key = f"tenants/{tenant_id}/chat-attachments/{conversation_id}/{date_path}/{file_id}_{filename}"

        # Upload to S3
        logger.info(f"Uploading chat attachment: {filename} ({content_type})")
        result = self.s3_storage.upload_file(
            file_content=file_content,
            key=s3_key,
            content_type=content_type,
            metadata={
                "tenant_id": str(tenant_id),
                "conversation_id": str(conversation_id),
                "original_filename": filename,
                "file_id": file_id,
            },
        )

        # Extract text from documents if applicable
        extracted_text = None
        if content_type in SUPPORTED_DOCUMENT_TYPES:
            try:
                extracted_text = await self._extract_text(
                    file_content=file_content if isinstance(file_content, bytes) else file_content.read(),
                    content_type=content_type,
                    filename=filename,
                )
                logger.info(f"Extracted {len(extracted_text)} characters from {filename}")
            except Exception as e:
                logger.warning(f"Failed to extract text from {filename}: {e}")

        # Generate presigned URL for download (valid for 1 hour)
        file_url = self.s3_storage.generate_presigned_url(s3_key, expiration=3600)

        # Generate thumbnail URL for images
        thumbnail_url = None
        if content_type in SUPPORTED_IMAGE_TYPES:
            thumbnail_url = file_url  # For now, use same URL. Can add thumbnail generation later

        # Build attachment metadata
        attachment = {
            "file_id": file_id,
            "file_name": filename,
            "file_type": content_type,
            "file_size": file_size,
            "file_url": result["url"],  # S3 URL
            "download_url": file_url,  # Presigned URL
            "thumbnail_url": thumbnail_url,
            "extracted_text": extracted_text,
            "s3_key": s3_key,
            "uploaded_at": timestamp.isoformat(),
        }

        logger.info(f"Successfully uploaded attachment: {file_id}")
        return attachment

    async def _extract_text(
        self,
        file_content: bytes,
        content_type: str,
        filename: str,
    ) -> str | None:
        """
        Extract text from document.

        Args:
            file_content: File content as bytes
            content_type: MIME type
            filename: Original filename

        Returns:
            str: Extracted text or None
        """
        try:
            # For plain text files, just decode
            if content_type == "text/plain":
                return file_content.decode("utf-8")

            # For markdown files
            if content_type == "text/markdown":
                return file_content.decode("utf-8")

            # For CSV files
            if content_type == "text/csv":
                return file_content.decode("utf-8")

            # For PDF files
            if content_type == "application/pdf":
                return self._extract_text_from_pdf(file_content)

            # For DOCX files
            if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_text_from_docx(file_content)

            logger.info(f"Text extraction for {content_type} not supported")
            return None

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return None

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        text_parts = []

        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        text_parts = []

        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise

    def validate_attachments(self, attachments: list[dict]) -> None:
        """
        Validate attachment list.

        Args:
            attachments: List of attachment metadata

        Raises:
            ValueError: If validation fails
        """
        if len(attachments) > MAX_FILES_PER_MESSAGE:
            raise ValueError(f"Too many attachments. Maximum {MAX_FILES_PER_MESSAGE} files per message.")

        total_size = sum(att.get("file_size", 0) for att in attachments)
        if total_size > MAX_FILE_SIZE * MAX_FILES_PER_MESSAGE:
            raise ValueError("Total attachment size exceeds maximum allowed size")

    def get_attachment_context(self, attachments: list[dict]) -> str:
        """
        Generate context string from attachments for LLM.

        Args:
            attachments: List of attachment metadata

        Returns:
            str: Context string
        """
        if not attachments:
            return ""

        context_parts = ["The user has attached the following files:"]

        for att in attachments:
            file_name = att.get("file_name", "unknown")
            file_type = att.get("file_type", "unknown")
            file_size = att.get("file_size", 0)

            context_parts.append(f"\n- {file_name} ({file_type}, {file_size} bytes)")

            # Include extracted text if available
            extracted_text = att.get("extracted_text")
            if extracted_text:
                # Truncate if too long
                max_text_length = 5000
                if len(extracted_text) > max_text_length:
                    extracted_text = extracted_text[:max_text_length] + "... (truncated)"
                context_parts.append(f"  Content:\n  {extracted_text}")

        return "\n".join(context_parts)

    def is_image_attachment(self, attachment: dict) -> bool:
        """Check if attachment is an image."""
        return attachment.get("file_type") in SUPPORTED_IMAGE_TYPES

    def get_image_attachments(self, attachments: list[dict]) -> list[dict]:
        """Get only image attachments from list."""
        return [att for att in attachments if self.is_image_attachment(att)]
