"""
Agent Context File Processor

Service for processing context files attached to agents.
Handles file validation, S3 upload, and text extraction.
"""

import csv
import io
import logging
import mimetypes
import uuid
from pathlib import Path
from typing import BinaryIO

import docx
import PyPDF2
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from src.models.agent import Agent
from src.models.agent_context_file import AgentContextFile
from src.services.storage.s3_storage import S3StorageService

logger = logging.getLogger(__name__)


class AgentContextFileProcessor:
    """Process context files for agents."""

    # File type configurations
    SUPPORTED_TYPES = {
        "application/pdf": {"ext": ".pdf", "max_size": 50 * 1024 * 1024},  # 50MB
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {
            "ext": ".docx",
            "max_size": 50 * 1024 * 1024,
        },
        "text/plain": {"ext": ".txt", "max_size": 10 * 1024 * 1024},  # 10MB
        "text/markdown": {"ext": ".md", "max_size": 10 * 1024 * 1024},
        "text/csv": {"ext": ".csv", "max_size": 10 * 1024 * 1024},
        "text/html": {"ext": ".html", "max_size": 10 * 1024 * 1024},
    }

    MAX_FILES_PER_AGENT = 100
    MAX_TOTAL_SIZE_PER_AGENT = 200 * 1024 * 1024  # 200MB

    def __init__(self, db: AsyncSession):
        """
        Initialize context file processor.

        Args:
            db: Async database session
        """
        self.db = db
        self.s3_storage = S3StorageService()

    def validate_file(self, filename: str, file_size: int, content_type: str | None = None) -> tuple[bool, str | None]:
        """
        Validate file before processing.

        Args:
            filename: Original filename
            file_size: File size in bytes
            content_type: MIME type (optional, will be guessed if not provided)

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Guess content type if not provided
        if not content_type:
            content_type, _ = mimetypes.guess_type(filename)

        if not content_type:
            return False, "Could not determine file type"

        # Check if file type is supported
        if content_type not in self.SUPPORTED_TYPES:
            supported_exts = [cfg["ext"] for cfg in self.SUPPORTED_TYPES.values()]
            return False, f"Unsupported file type. Supported types: {', '.join(supported_exts)}"

        # Check file size
        max_size = self.SUPPORTED_TYPES[content_type]["max_size"]
        if file_size > max_size:
            max_size_mb = max_size / (1024 * 1024)
            return False, f"File size exceeds maximum of {max_size_mb}MB"

        return True, None

    async def validate_agent_limits(self, agent: Agent, new_file_size: int) -> tuple[bool, str | None]:
        """
        Validate agent file limits.

        Args:
            agent: Agent instance
            new_file_size: Size of new file to add

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Use an aggregate query — avoids lazy-loading the relationship in an async session
        result = await self.db.execute(
            select(func.count(AgentContextFile.id), func.coalesce(func.sum(AgentContextFile.file_size), 0)).filter(
                AgentContextFile.agent_id == agent.id
            )
        )
        current_file_count, current_total_size = result.one()

        if current_file_count >= self.MAX_FILES_PER_AGENT:
            return False, f"Maximum of {self.MAX_FILES_PER_AGENT} files per agent"

        if current_total_size + new_file_size > self.MAX_TOTAL_SIZE_PER_AGENT:
            max_size_mb = self.MAX_TOTAL_SIZE_PER_AGENT / (1024 * 1024)
            return False, f"Total file size would exceed maximum of {max_size_mb}MB"

        return True, None

    async def process_file(
        self, agent: Agent, file: BinaryIO, filename: str, content_type: str | None = None
    ) -> AgentContextFile:
        """
        Process and store a context file for an agent.

        Args:
            agent: Agent instance
            file: File object
            filename: Original filename
            content_type: MIME type (optional)

        Returns:
            AgentContextFile instance

        Raises:
            ValueError: If validation fails
            Exception: If processing fails
        """
        # Read file content
        file_content = file.read()
        file_size = len(file_content)

        # Guess content type if not provided
        if not content_type:
            content_type, _ = mimetypes.guess_type(filename)

        # Validate file
        is_valid, error = self.validate_file(filename, file_size, content_type)
        if not is_valid:
            raise ValueError(error)

        # Validate agent limits
        is_valid, error = await self.validate_agent_limits(agent, file_size)
        if not is_valid:
            raise ValueError(error)

        try:
            # Generate S3 key
            file_uuid = uuid.uuid4()
            file_ext = Path(filename).suffix
            s3_key = f"agent-context/agent-{agent.id}/context-files/file-{file_uuid}{file_ext}"

            # Upload to S3
            logger.info(f"Uploading file {filename} to S3: {s3_key}")
            await self.s3_storage.upload_file_content(
                file_content=file_content,
                key=s3_key,
                content_type=content_type,
                metadata={"agent_id": str(agent.id), "original_filename": filename, "tenant_id": str(agent.tenant_id)},
            )

            # Get S3 bucket from config
            s3_bucket = self.s3_storage.bucket_name

            # Get current file count for display_order (async query — avoids lazy-load)
            count_result = await self.db.execute(
                select(func.count(AgentContextFile.id)).filter(AgentContextFile.agent_id == agent.id)
            )
            current_count = count_result.scalar() or 0

            # Create database record
            context_file = AgentContextFile(
                agent_id=agent.id,
                tenant_id=agent.tenant_id,
                filename=filename,
                file_type=content_type,
                file_size=file_size,
                s3_key=s3_key,
                s3_bucket=s3_bucket,
                extraction_status="PENDING",
                display_order=current_count,
            )

            self.db.add(context_file)
            await self.db.flush()  # Get the ID

            # Extract text asynchronously
            try:
                logger.info(f"Extracting text from file {filename}")
                extracted_text = await self._extract_text(file_content, content_type, filename)

                context_file.extracted_text = extracted_text
                context_file.extraction_status = "COMPLETED"
                logger.info(f"Successfully extracted {len(extracted_text)} characters from {filename}")

            except Exception as e:
                logger.error(f"Error extracting text from {filename}: {e}")
                context_file.extraction_status = "FAILED"
                context_file.extraction_error = str(e)

            await self.db.commit()

            logger.info(f"Successfully processed context file {filename} for agent {agent.agent_name}")
            return context_file

        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error processing file {filename}: {e}")
            raise

    async def _extract_text(self, file_content: bytes, content_type: str, filename: str) -> str:
        """
        Extract text from file content.

        Args:
            file_content: File content as bytes
            content_type: MIME type
            filename: Original filename

        Returns:
            Extracted text
        """
        try:
            if content_type == "application/pdf":
                return self._extract_text_from_pdf(file_content)

            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_text_from_docx(file_content)

            elif content_type in ["text/plain", "text/markdown", "text/html"]:
                return file_content.decode("utf-8", errors="ignore")

            elif content_type == "text/csv":
                return self._extract_text_from_csv(file_content)

            else:
                raise ValueError(f"Unsupported content type for text extraction: {content_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        text_parts = []

        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        text_parts = []

        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise

    def _extract_text_from_csv(self, file_content: bytes) -> str:
        """Extract text from CSV file."""
        text_parts = []

        try:
            csv_text = file_content.decode("utf-8", errors="ignore")
            csv_file = io.StringIO(csv_text)
            csv_reader = csv.reader(csv_file)

            for row_num, row in enumerate(csv_reader):
                row_text = " | ".join(str(cell).strip() for cell in row)
                if row_text.strip():
                    text_parts.append(f"Row {row_num + 1}: {row_text}")

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            raise

    async def delete_file(self, context_file: AgentContextFile) -> None:
        """
        Delete a context file.

        Args:
            context_file: AgentContextFile instance
        """
        try:
            # Delete from S3
            logger.info(f"Deleting file from S3: {context_file.s3_key}")
            self.s3_storage.delete_file(context_file.s3_key)

            # Delete from database
            await self.db.delete(context_file)
            await self.db.commit()

            logger.info(f"Successfully deleted context file {context_file.filename}")

        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error deleting file {context_file.filename}: {e}")
            raise

    async def get_download_url(self, context_file: AgentContextFile) -> str:
        """
        Get a presigned download URL for a context file.

        Args:
            context_file: AgentContextFile instance

        Returns:
            Presigned download URL
        """
        try:
            url = await self.s3_storage.generate_presigned_url(
                key=context_file.s3_key,
                expiration=3600,  # 1 hour
            )
            return url

        except Exception as e:
            logger.error(f"Error generating download URL for {context_file.filename}: {e}")
            raise

    def get_context_files_text(self, agent: Agent) -> str:
        """
        Get combined text from all context files for an agent.

        Args:
            agent: Agent instance

        Returns:
            Combined text from all context files
        """
        text_parts = []

        for context_file in agent.context_files:
            if context_file.is_extraction_complete and context_file.extracted_text:
                text_parts.append(f"=== Context File: {context_file.filename} ===\n\n{context_file.extracted_text}\n\n")

        return "\n".join(text_parts)
